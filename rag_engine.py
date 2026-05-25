import os
import re
import base64
import io
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

class TachografRAG:
    def __init__(self, persist_directory="./chroma_db_prawdziwa"):
        self.persist_directory = persist_directory
        self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        self.vectorstore = None
        self.chain = None
        
    def initialize_database(self, processed_chunks):
        print("Inicjalizacja bazy wektorowej ChromaDB...")
        documents = [Document(page_content=c["text"], metadata=c["metadata"]) for c in processed_chunks]
        self.vectorstore = Chroma.from_documents(documents=documents, embedding=self.embeddings, persist_directory=self.persist_directory)
        self._build_rag_chain()

    def load_existing_database(self):
        if os.path.exists(self.persist_directory):
            print("Ładowanie istniejącej bazy ChromaDB...")
            self.vectorstore = Chroma(persist_directory=self.persist_directory, embedding_function=self.embeddings)
            self._build_rag_chain()
            return True
        return False

    def _build_rag_chain(self):
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 20})
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.0)
        
        template = """Jesteś wybitnym ekspertem ds. europejskiego prawa transportowego i głównym doradcą DGSA.
Twoim zadaniem jest odpowiadanie na pytania użytkowników WYŁĄCZNIE na podstawie dostarczonego poniżej KONTEKSTU.

ZASADY BEZWZGLĘDNE:
1. BRAK ZGADYWANIA: Jeśli w kontekście NIE MA odpowiedzi, powiedz wprost o braku artykułu.
2. PRECYZJA: Opieraj się na konkretnych artykułach i kodach.
3. CYTOWANIE: Zawsze dodaj: [ŹRÓDŁO: nazwa dokumentu z metadanych].
4. WIELOJĘZYCZNOŚĆ: ZAWSZE odpowiadaj dokładnie w tym samym języku, w którym użytkownik zadał pytanie.

KONTEKST PRAWNY:
{context}

ZAPYTANIE UŻYTKOWNIKA:
{question}
ODPOWIEDŹ:"""
        
        prompt = ChatPromptTemplate.from_template(template)
        
        def format_docs(docs): 
            slownik_zrodel = {
                "data/t1.pdf": "Umowa ADR 2023 (Tom I)",
                "data/t2.pdf": "Umowa ADR 2023 (Tom II)",
                "data/r1.pdf": "Rozporządzenie (WE) nr 561/2006 (Czas Pracy)",
                "data/r2.pdf": "Rozporządzenie (UE) nr 165/2014 (Tachografy)",
                "data/r3.pdf": "Pakiet Mobilności (Rozporządzenie 1054/2020)",
                "data/r4.pdf": "Europejskie Prawo Transportowe"
            }
            wyniki = []
            for doc in docs:
                surowe_zrodlo = doc.metadata.get('source', 'Nieznane źródło')
                ladne_zrodlo = slownik_zrodel.get(surowe_zrodlo, surowe_zrodlo)
                wyniki.append(f"Tekst: {doc.page_content}\nŹródło: {ladne_zrodlo}")
            return "\n\n=== KOLEJNY PRZEPIS ===\n\n".join(wyniki)
        
        self.chain = (
            {"context": retriever | format_docs, "question": RunnablePassthrough()}
            | prompt | llm | StrOutputParser()
        )

    def ask(self, question: str) -> str:
        if not self.chain: raise ValueError("Brak bazy!")
        return self.chain.invoke(question)

    # --- NOWE FUNKCJE MULTIMEDIALNE ---
    
    def read_image(self, image_bytes):
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        chat = ChatOpenAI(model="gpt-4o-mini", max_tokens=1000)
        msg = chat.invoke([
            HumanMessage(content=[
                {"type": "text", "text": "Odczytaj dokładnie dane z dokumentu. Zwróć czysty tekst."},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
            ])
        ])
        return msg.content

    def read_pdf(self, file_obj):
        import PyPDF2
        reader = PyPDF2.PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            extr = page.extract_text()
            if extr: text += extr + "\n"
        return text

    def read_docx(self, file_obj):
        import docx
        doc = docx.Document(io.BytesIO(file_obj.read()))
        return "\n".join([p.text for p in doc.paragraphs])

    def transcribe_audio(self, audio_file_obj):
        from openai import OpenAI
        client = OpenAI()
        audio_file_obj.name = "audio.wav" # Wymóg API OpenAI
        transcript = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file_obj
        )
        return transcript.text

    # --- FUNKCJE PRAWNE ---

    def calculate_penalty(self, violation_description: str) -> str:
        from openai import OpenAI
        results = self.vectorstore.similarity_search(violation_description, k=6)
        context_text = "\n\n---\n\n".join([f"Tekst: {doc.page_content}\nŹródło: {doc.metadata.get('source', 'Brak źródła')}" for doc in results])
        system_prompt = """
        Jesteś Inspektorem ITD/BAG i ekspertem ds. ryzyka finansowego w transporcie.
        Wykonaj wycenę naruszenia na podstawie kontekstu.
        BARDZO WAŻNE: Twoja ostateczna odpowiedź MUSI być w tym samym języku, w którym napisany jest opis naruszenia od użytkownika.
        """
        try:
            client = OpenAI()
            response = client.chat.completions.create(
                model="gpt-4o", 
                temperature=0.0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Kontekst:\n{context_text}\n\nNaruszenie: {violation_description}"}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Błąd modułu finansowego: {str(e)}"

    def generate_defense_statement(self, incident_description: str, target_language: str = "Polski") -> str:
        from openai import OpenAI
        results = self.vectorstore.similarity_search(incident_description, k=8)
        context_text = "\n\n---\n\n".join([f"Tekst: {doc.page_content}\nŹródło: {doc.metadata.get('source', 'Brak źródła')}" for doc in results])
        system_prompt = f"""
        Jesteś wybitnym Adwokatem specjalizującym się w prawie transportowym.
        Twoim zadaniem jest wygenerowanie profesjonalnego oświadczenia dla kierowcy (np. na podstawie Art. 12 Rozporządzenia 561/2006).
        
        WYTYCZNE:
        1. JĘZYK OŚWIADCZENIA: Sam tekst pisma obronnego MUSI być przetłumaczony na język: {target_language}.
        2. Napisz w tonie oficjalnego dokumentu dla służb kontrolnych (BAG, DREAL, Guardia Civil, ITD).
        3. Powołaj się na konkretne artykuły z dostarczonego kontekstu.
        4. Zwracaj gotowy tekst pisma, luki oznacz jako [Imię i Nazwisko], [Data].
        5. Pod pismem dodaj krótką notatkę dla kierowcy W JĘZYKU POLSKIM, wyjaśniającą, dlaczego taka linia obrony została przyjęta.
        """
        try:
            client = OpenAI()
            response = client.chat.completions.create(
                model="gpt-4o", 
                temperature=0.2,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Przepisy (Kontekst):\n{context_text}\n\nOpis sytuacji od kierowcy: {incident_description}"}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Błąd modułu prawniczego: {str(e)}"